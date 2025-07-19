import os
import json
import shutil
import zipfile
import logging
from datetime import datetime
from flask import Flask, render_template, request, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from pydub import AudioSegment
from docx import Document

# Configure logging
logging.basicConfig(level=logging.DEBUG)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)

# Flask App Initialization
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "forensic-clueword-extractor-secret")

# Database configuration  
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL") or "sqlite:///forensic_sessions.db"
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db.init_app(app)

# Database Model
class ForensicSession(db.Model):
    __tablename__ = 'forensic_sessions'
    
    id = db.Column(db.Integer, primary_key=True)
    session_name = db.Column(db.String(200), nullable=False)
    case_number = db.Column(db.String(100))
    police_station = db.Column(db.String(200))
    district = db.Column(db.String(100))
    cr_number = db.Column(db.String(100))
    speaker_name = db.Column(db.String(200))
    
    # Audio file information
    question_filename = db.Column(db.String(500))
    control_filename = db.Column(db.String(500))
    question_file_path = db.Column(db.String(500))
    control_file_path = db.Column(db.String(500))
    
    # Session data
    annotations_data = db.Column(db.Text)  # JSON string of all annotations
    bandpass_enabled = db.Column(db.Boolean, default=False)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    def set_annotations(self, annotations):
        """Store annotations as JSON string"""
        self.annotations_data = json.dumps(annotations)
    
    def get_annotations(self):
        """Retrieve annotations from JSON string"""
        if self.annotations_data:
            return json.loads(self.annotations_data)
        return {"question": [], "control": []}
    
    def to_dict(self):
        """Convert session to dictionary for JSON serialization"""
        return {
            'id': self.id,
            'session_name': self.session_name,
            'case_number': self.case_number,
            'police_station': self.police_station,
            'district': self.district,
            'cr_number': self.cr_number,
            'speaker_name': self.speaker_name,
            'question_filename': self.question_filename,
            'control_filename': self.control_filename,
            'annotations': self.get_annotations(),
            'bandpass_enabled': self.bandpass_enabled,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

# Initialize database tables (only if they don't exist)
with app.app_context():
    try:
        db.create_all()
    except Exception as e:
        app.logger.info(f"Tables might already exist: {e}")
        pass

# Configuration
# Define paths for file storage
UPLOAD_FOLDER = 'uploads'
STANDARDIZED_FOLDER = os.path.join('static', 'temp_standardized')
OUTPUT_FOLDER = 'output'

def clear_directory(folder_path):
    """Helper function to clean up directories"""
    if not os.path.exists(folder_path):
        os.makedirs(folder_path, exist_ok=True)
        return
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            app.logger.error(f"Error clearing {file_path}: {str(e)}")

@app.route('/')
def index():
    """Renders the main HTML page and clears old session data."""
    try:
        # Clean up directories for a fresh session
        for folder in [UPLOAD_FOLDER, STANDARDIZED_FOLDER, OUTPUT_FOLDER]:
            clear_directory(folder)
        
        # Clear the root-level zip file if it exists
        if os.path.exists('clueword_analysis.zip'):
            os.remove('clueword_analysis.zip')
            
        return render_template('index.html')
    except Exception as e:
        app.logger.error(f"Error in index route: {str(e)}")
        return render_template('index.html')

@app.route('/standardize', methods=['POST'])
def standardize_audio():
    """
    Receives an audio file, converts it to a standard format,
    and returns the path to the new file.
    """
    try:
        if 'audio_file' not in request.files:
            return jsonify({"error": "No audio file provided."}), 400
        
        file = request.files['audio_file']
        panel_type = request.form.get('type')  # 'question' or 'control'
        
        if not file.filename:
            return jsonify({"error": "No file selected."}), 400
            
        if not panel_type or panel_type not in ['question', 'control']:
            return jsonify({"error": "Invalid panel type."}), 400
        
        # Ensure upload folder exists
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        os.makedirs(STANDARDIZED_FOLDER, exist_ok=True)
        
        original_filename = f"{panel_type}_original"
        original_path = os.path.join(UPLOAD_FOLDER, original_filename)
        file.save(original_path)
        
        # Load and standardize audio
        audio = AudioSegment.from_file(original_path)
        standardized_audio = audio.set_frame_rate(44100).set_channels(1).set_sample_width(2)
        
        standardized_filename = f"{panel_type}_standardized.wav"
        standardized_path = os.path.join(STANDARDIZED_FOLDER, standardized_filename)
        standardized_audio.export(standardized_path, format="wav")
        
        return jsonify({
            "success": True,
            "url": f"/static/temp_standardized/{standardized_filename}",
            "original_filename": file.filename
        })

    except Exception as e:
        app.logger.error(f"Error in standardize_audio: {str(e)}")
        return jsonify({"error": f"Audio processing failed: {str(e)}"}), 500

@app.route('/process', methods=['POST'])
def process_audio():
    """
    Receives annotation data and generates the final cluewords package.
    """
    try:
        annotations_data = request.form.get('annotations')
        if not annotations_data:
            return jsonify({"error": "No annotations provided."}), 400
            
        annotations = json.loads(annotations_data)
        q_annotations = annotations.get('question', [])
        c_annotations = annotations.get('control', [])
        
        if not q_annotations or not c_annotations:
            return jsonify({"error": "Both question and control annotations are required."}), 400
        
        q_original_filename = request.form.get('question_original_filename')
        c_original_filename = request.form.get('control_original_filename')
        enable_bandpass = request.form.get('enable_bandpass', 'true').lower() == 'true'
        
        # Get case information
        case_info = {
            'case_number': request.form.get('case_number', ''),
            'police_station': request.form.get('police_station', ''),
            'district': request.form.get('district', ''),
            'cr_adr_number': request.form.get('cr_adr_number', ''),
            'speaker_name': request.form.get('speaker_name', '')
        }
        
        if not q_original_filename or not c_original_filename:
            return jsonify({"error": "Original filenames are required."}), 400

        # Load original audio files
        q_audio_path = os.path.join(UPLOAD_FOLDER, "question_original")
        c_audio_path = os.path.join(UPLOAD_FOLDER, "control_original")
        
        if not os.path.exists(q_audio_path) or not os.path.exists(c_audio_path):
            return jsonify({"error": "Original audio files not found."}), 400
            
        q_audio = AudioSegment.from_file(q_audio_path)
        c_audio = AudioSegment.from_file(c_audio_path)

        report_data = []
        matches_found = 0
        
        # Create a map of control annotations for quick lookup
        control_map = {ann['label'].lower().strip(): ann for ann in c_annotations}

        # Clear output directory
        clear_directory(OUTPUT_FOLDER)

        # Process question annotations and find matches
        for q_ann in q_annotations:
            q_label_clean = q_ann['label'].lower().strip()
            if q_label_clean in control_map:
                matches_found += 1
                c_ann = control_map[q_label_clean]
                
                # Create safe directory name
                safe_label = "".join(c for c in q_ann['label'] if c.isalnum() or c in (' ', '_')).rstrip()
                if not safe_label:
                    safe_label = f"clueword_{matches_found}"
                    
                clueword_dir = os.path.join(OUTPUT_FOLDER, safe_label)
                os.makedirs(clueword_dir, exist_ok=True)

                # Extract question segment
                q_start_ms = float(q_ann['start']) * 1000
                q_end_ms = float(q_ann['end']) * 1000
                q_seg = q_audio[q_start_ms:q_end_ms].set_channels(1).set_frame_rate(44100)
                q_seg.export(os.path.join(clueword_dir, "question.wav"), format="wav")
                
                # Apply bandpass filter to question segment if enabled (400Hz - 4000Hz)
                if enable_bandpass:
                    q_seg_bpf = apply_bandpass_filter(q_seg, 400, 4000)
                    q_seg_bpf.export(os.path.join(clueword_dir, "bpf_question.wav"), format="wav")
                
                report_data.append(["Question", q_ann['label'], q_start_ms, q_end_ms, q_end_ms - q_start_ms])

                # Extract control segment
                c_start_ms = float(c_ann['start']) * 1000
                c_end_ms = float(c_ann['end']) * 1000
                c_seg = c_audio[c_start_ms:c_end_ms].set_channels(1).set_frame_rate(44100)
                c_seg.export(os.path.join(clueword_dir, "control.wav"), format="wav")
                
                # Apply bandpass filter to control segment if enabled (400Hz - 4000Hz)
                if enable_bandpass:
                    c_seg_bpf = apply_bandpass_filter(c_seg, 400, 4000)
                    c_seg_bpf.export(os.path.join(clueword_dir, "bpf_control.wav"), format="wav")
                
                report_data.append(["Control", c_ann['label'], c_start_ms, c_end_ms, c_end_ms - c_start_ms])
        
        if matches_found == 0:
            return jsonify({"error": "No matching annotations found between question and control files."}), 400

        # Create analysis report
        create_report(report_data, OUTPUT_FOLDER, q_original_filename, c_original_filename, matches_found, enable_bandpass, case_info)

        # Create ZIP file
        zip_path = 'clueword_analysis.zip'
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(OUTPUT_FOLDER):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_name = os.path.relpath(file_path, OUTPUT_FOLDER)
                    zipf.write(file_path, arc_name)
        
        return send_file(zip_path, as_attachment=True, download_name='clueword_analysis.zip')

    except json.JSONDecodeError:
        return jsonify({"error": "Invalid JSON in annotations data."}), 400
    except Exception as e:
        app.logger.error(f"Error in process_audio: {str(e)}")
        return jsonify({"error": "An internal server error occurred during processing."}), 500

def create_report(data, output_dir, q_filename, c_filename, matches_count, enable_bandpass=True, case_info=None):
    """Generates a comprehensive .docx report."""
    try:
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Clueword Sheet', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page numbers
        from docx.oxml import parse_xml
        section = doc.sections[0]
        header = section.header
        footer = section.footer
        
        # Add page number to footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "
        run = footer_para.runs[0]
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fldChar1)
        
        instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
        run._r.append(instrText)
        
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fldChar2)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case Information Header (Bold format as requested)
        if case_info:
            case_header = doc.add_heading('Header Title:', level=1)
            case_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Create case information table
            case_table = doc.add_table(rows=2, cols=2)
            case_table.style = 'Table Grid'
            
            # Row 1
            case_table.cell(0, 0).text = "Case No."
            case_table.cell(0, 1).text = case_info.get('case_number', 'N/A')
            case_table.cell(1, 0).text = "Police Station Name"
            case_table.cell(1, 1).text = case_info.get('police_station', 'N/A')
            
            # Add more rows for district and CR/ADR
            row = case_table.add_row()
            row.cells[0].text = "District"
            row.cells[1].text = case_info.get('district', 'N/A')
            
            row = case_table.add_row()
            row.cells[0].text = "C.R./A.D.R. No."
            row.cells[1].text = case_info.get('cr_adr_number', 'N/A')
            
            row = case_table.add_row()
            row.cells[0].text = "Speaker Name"
            row.cells[1].text = case_info.get('speaker_name', 'N/A')
            
            # Make labels bold
            for row in case_table.rows:
                for i, cell in enumerate(row.cells):
                    if i == 0:  # First column (labels)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            doc.add_paragraph()
        
        # Report metadata
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Matching Cluewords Found: {matches_count}")
        doc.add_paragraph()
        
        # Add summary
        doc.add_heading('Analysis Summary', level=1)
        doc.add_paragraph(f"This forensic analysis identified {matches_count} matching cluewords between the question and control audio recordings. Each matching clueword has been extracted as separate audio segments for further analysis.")
        doc.add_paragraph()
        
        # Add detailed table with new format
        doc.add_heading('Detailed Analysis', level=1)
        table = doc.add_table(rows=1, cols=7, style='Table Grid')
        table.autofit = False
        
        # Table headers - new format
        hdr_cells = table.rows[0].cells
        headers = [
            f'Question File: {q_filename}', 
            '', 
            '', 
            '', 
            f'Control File: {c_filename}', 
            '', 
            ''
        ]
        
        # Merge cells for file headers
        hdr_cells[0].merge(hdr_cells[3])
        hdr_cells[4].merge(hdr_cells[6])
        
        for i, header in enumerate([f'Question File: {q_filename}', f'Control File: {c_filename}']):
            cell_idx = 0 if i == 0 else 4
            hdr_cells[cell_idx].text = header
            for paragraph in hdr_cells[cell_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add subheaders row
        subhdr_row = table.add_row()
        subhdr_cells = subhdr_row.cells
        subheaders = [
            'Clueword', 
            'Start (HH:MM:SS:MS)', 
            'End (HH:MM:SS:MS)', 
            'Duration (ms)',
            'Start (HH:MM:SS:MS)', 
            'End (HH:MM:SS:MS)', 
            'Duration (ms)'
        ]
        
        for i, header in enumerate(subheaders):
            subhdr_cells[i].text = header
            for paragraph in subhdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Group data by matching cluewords
        clueword_matches = {}
        for item in data:
            source, label, start_ms, end_ms, duration_ms = item
            if label not in clueword_matches:
                clueword_matches[label] = {'question': None, 'control': None}
            
            clueword_matches[label][source.lower()] = {
                'start_ms': start_ms,
                'end_ms': end_ms,
                'duration_ms': duration_ms
            }
        
        # Add data rows - one row per matching clueword pair
        for label, match_data in clueword_matches.items():
            if match_data['question'] and match_data['control']:
                row_cells = table.add_row().cells
                
                # Question data
                q_data = match_data['question']
                row_cells[0].text = label
                row_cells[1].text = format_time_hhmmssms(q_data['start_ms'])
                row_cells[2].text = format_time_hhmmssms(q_data['end_ms'])
                row_cells[3].text = f"{q_data['duration_ms']:.0f}"
                
                # Control data (no clueword column - it's already shown in the first column)
                c_data = match_data['control']
                row_cells[4].text = format_time_hhmmssms(c_data['start_ms'])
                row_cells[5].text = format_time_hhmmssms(c_data['end_ms'])
                row_cells[6].text = f"{c_data['duration_ms']:.0f}"
        
        doc.add_paragraph()
        doc.add_heading('Notes', level=1)
        doc.add_paragraph("• All audio segments have been standardized to 44100Hz, mono, 16-bit format")
        doc.add_paragraph("• Matching is performed using case-insensitive label comparison")
        doc.add_paragraph("• Time values are referenced to the original audio files")
        if enable_bandpass:
            doc.add_paragraph("• Each clueword directory contains 4 files: question.wav, control.wav, bpf_question.wav, bpf_control.wav")
            doc.add_paragraph("• BPF (Bandpass Filtered) files use 400Hz-4000Hz frequency range for enhanced voice clarity")
        else:
            doc.add_paragraph("• Each clueword directory contains 2 files: question.wav, control.wav")
        doc.add_paragraph("• Duration values are shown in milliseconds")
        
        doc.save(os.path.join(output_dir, "analysis_report.docx"))
        
    except Exception as e:
        app.logger.error(f"Error creating report: {str(e)}")
        # Create a simple fallback report
        with open(os.path.join(output_dir, "analysis_report.txt"), 'w') as f:
            f.write(f"Forensic Clueword Analysis Report\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Question File: {q_filename}\n")
            f.write(f"Control File: {c_filename}\n")
            f.write(f"Matches Found: {matches_count}\n\n")
            f.write("Detailed Analysis:\n")
            for item in data:
                f.write(f"{item}\n")

# Session Management Routes
@app.route('/api/sessions', methods=['GET'])
def get_sessions():
    """Get all saved forensic sessions"""
    try:
        sessions = ForensicSession.query.order_by(ForensicSession.updated_at.desc()).all()
        return jsonify([session.to_dict() for session in sessions])
    except Exception as e:
        app.logger.error(f"Error fetching sessions: {str(e)}")
        return jsonify({'error': 'Failed to fetch sessions'}), 500

@app.route('/api/sessions', methods=['POST'])
def save_session():
    """Save a forensic session"""
    try:
        data = request.get_json()
        
        session_id = data.get('session_id')
        if session_id:
            # Update existing session
            session = ForensicSession.query.get(session_id)
            if not session:
                return jsonify({'error': 'Session not found'}), 404
        else:
            # Create new session
            session = ForensicSession()
        
        # Update session data
        session.session_name = data.get('session_name', '')
        session.case_number = data.get('case_number', '')
        session.police_station = data.get('police_station', '')
        session.district = data.get('district', '')
        session.cr_number = data.get('cr_number', '')
        session.speaker_name = data.get('speaker_name', '')
        session.question_filename = data.get('question_filename', '')
        session.control_filename = data.get('control_filename', '')
        session.question_file_path = data.get('question_file_path', '')
        session.control_file_path = data.get('control_file_path', '')
        session.bandpass_enabled = data.get('bandpass_enabled', False)
        
        # Save annotations
        annotations = data.get('annotations', {"question": [], "control": []})
        session.set_annotations(annotations)
        
        db.session.add(session)
        db.session.commit()
        
        return jsonify(session.to_dict())
        
    except Exception as e:
        app.logger.error(f"Error saving session: {str(e)}")
        db.session.rollback()
        return jsonify({'error': 'Failed to save session'}), 500

@app.route('/api/sessions/<int:session_id>', methods=['GET'])
def load_session(session_id):
    """Load a specific forensic session"""
    try:
        session = ForensicSession.query.get(session_id)
        if not session:
            return jsonify({'error': 'Session not found'}), 404
        
        return jsonify(session.to_dict())
        
    except Exception as e:
        app.logger.error(f"Error loading session: {str(e)}")
        return jsonify({'error': 'Failed to load session'}), 500

@app.route('/api/sessions/<int:session_id>', methods=['DELETE'])
def delete_session(session_id):
    """Delete a forensic session"""
    try:
        session = ForensicSession.query.get(session_id)
        if not session:
            return jsonify({'error': 'Session not found'}), 404
        
        db.session.delete(session)
        db.session.commit()
        
        return jsonify({'message': 'Session deleted successfully'})
        
    except Exception as e:
        app.logger.error(f"Error deleting session: {str(e)}")
        db.session.rollback()
        return jsonify({'error': 'Failed to delete session'}), 500

def apply_bandpass_filter(audio_segment, low_freq, high_freq):
    """Apply bandpass filter using FFmpeg"""
    import tempfile
    import os
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_input:
            audio_segment.export(temp_input.name, format='wav')
            temp_input_path = temp_input.name
        
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_output:
            temp_output_path = temp_output.name
        
        # Apply bandpass filter using FFmpeg
        # highpass=400Hz, lowpass=4000Hz
        filter_cmd = f"highpass=f={low_freq},lowpass=f={high_freq}"
        
        import subprocess
        cmd = [
            'ffmpeg', '-y', '-i', temp_input_path,
            '-af', filter_cmd,
            '-acodec', 'pcm_s16le',
            temp_output_path
        ]
        
        subprocess.run(cmd, check=True, capture_output=True)
        
        # Load filtered audio
        filtered_audio = AudioSegment.from_wav(temp_output_path)
        
        # Clean up temp files
        os.unlink(temp_input_path)
        os.unlink(temp_output_path)
        
        return filtered_audio
        
    except Exception as e:
        app.logger.warning(f"Bandpass filter failed: {str(e)}, returning original audio")
        return audio_segment

def format_time_hhmmssms(milliseconds):
    """Convert milliseconds to HH:MM:SS:MS format"""
    total_seconds = milliseconds / 1000
    hours = int(total_seconds // 3600)
    minutes = int((total_seconds % 3600) // 60)
    seconds = int(total_seconds % 60)
    ms = int(milliseconds % 1000)
    
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}:{ms:03d}"

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
